from django.contrib.auth.views import LoginView
from django.contrib.auth import authenticate, login
from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponseRedirect, JsonResponse, HttpResponse
from django.urls import reverse_lazy,  reverse
from django.contrib import messages
from django.views.generic import TemplateView, View, CreateView,  ListView
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth import get_user_model
from django.contrib.auth.tokens import default_token_generator
from django.contrib.auth.forms import PasswordResetForm, SetPasswordForm
from .forms import NewResetPasswordForm
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.utils.encoding import force_bytes, force_str
from django.core.mail import send_mail
from django.conf import settings
from django.template.loader import render_to_string
from .forms import UserLoginForm
from django.contrib.auth import logout
from django.shortcuts import redirect
from django.db import IntegrityError
from django_tenants.utils import tenant_context
from django.utils.text import slugify
from django import forms
from django.utils import timezone
from django.contrib.auth.views import PasswordResetView, PasswordResetConfirmView
from django.db.models import Q
from django.contrib.sessions.models import Session
from django.utils.decorators import method_decorator
from django.contrib.auth.decorators import login_required
import json
from io import BytesIO
import os


# Imports para exportação
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
import io
from PIL import Image as PILImage
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# Imports para DOC
try:
    from docx import Document
    from docx.shared import Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


from .forms import CustomPasswordResetForm, UsuarioClienteForm, ClienteSystemSettingsForm
from .models import  Prova,  Prova, QuestaoProva ,UsuarioCliente, ClienteSystemSettings, SessaoUsuarioCliente
from .services import QuestaoService


class ClienteUserLoginView(LoginView):
    template_name = "clienteLogin.html"
    success_url = reverse_lazy('LSCliente:clientehome')
    form_class = UserLoginForm

    def get_form_kwargs(self):
        # Adiciona o request aos kwargs do formulário
        kwargs = super().get_form_kwargs()
        kwargs['request'] = self.request
        return kwargs
    
    def form_valid(self, form):
        """Método que é chamado quando o formulário é válido"""
        user = form.get_user()
        # Efetuar login manualmente para garantir que a sessão seja corretamente inicializada
        login(self.request, user)
        
        # Registrar o login no console para debug
        print(f"[LOGIN DEBUG] Usuário {user.email} autenticado com sucesso")
        print(f"[LOGIN DEBUG] Sessão: {self.request.session.session_key}")
        
        # Adicionar marcador de sessão para verificar no middleware
        self.request.session['authenticated_tenant'] = self.request.tenant.schema_name
        self.request.session.save()  # Forçar salvamento da sessão
        
        return HttpResponseRedirect(self.get_success_url())
    
    def post(self, request, *args, **kwargs):
        # Usa o método get_form_kwargs para garantir que o request seja passado
        form_kwargs = self.get_form_kwargs()
        form = self.form_class(**form_kwargs)
        
        if form.is_valid():
            return self.form_valid(form)
        return self.form_invalid(form)

    def get_success_url(self):
        next_url = self.request.GET.get('next')
        print(f"[DEBUG] ClienteUserLoginView: next_url={next_url}")
        
        if next_url:
            # Se for /home/, substituir pelo caminho correto
            if next_url == '/home/':
                return next_url  # Manter o caminho como está para evitar problemas de resolução
            return next_url
        
        default_url = reverse('LSCliente:clientehome')
        print(f"[DEBUG] Usando URL padrão: {default_url}")
        return default_url

class ClienteHomeView(LoginRequiredMixin, TemplateView):
    template_name = 'cliente_index.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["title"] = "Home"
        # Adicionar a lista de usuários ao contexto
        from LSCliente.models import UsuarioCliente
        context["usuarios"] = UsuarioCliente.objects.all().order_by('nome')
        context["tenant_name"] = self.request.tenant.nome
        context["tenant_schema"] = self.request.tenant.schema_name
        return context

class TenantPasswordResetView(PasswordResetView):
    """
    View para reset de senha que é consciente do tenant atual.
    """
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        # Passa o request para o formulário
        kwargs['request'] = self.request
        return kwargs

class TenantAwarePasswordResetConfirmView(View):
    """
    Uma versão completamente personalizada do PasswordResetConfirmView 
    que é consciente do sistema multi-tenant.
    """
    template_name = 'cliente_password_reset/password_reset_senha_nova_form.html'
    success_url = reverse_lazy('LSCliente:clientepassword_reset_complete')
    
    def get(self, request, *args, **kwargs):
        """
        Exibe o formulário de redefinição de senha e valida o token
        """
        # Debug
        print(f"[DEBUG] TenantAwarePasswordResetConfirmView GET chamado no tenant: {request.tenant.schema_name}")
        print(f"[DEBUG] uidb64: {kwargs.get('uidb64')}")
        print(f"[DEBUG] token: {kwargs.get('token')}")
        
        context = self._get_context_with_user(request, **kwargs)
        
        if context['validlink']:
            form = NewResetPasswordForm(user=context['user'])
            context['form'] = form
        
        return render(request, self.template_name, context)
    
    def post(self, request, *args, **kwargs):
        """
        Processa o formulário de redefinição de senha
        """
        # Debug
        print(f"[DEBUG] TenantAwarePasswordResetConfirmView POST chamado no tenant: {request.tenant.schema_name}")
        
        context = self._get_context_with_user(request, **kwargs)
        
        if not context['validlink']:
            return render(request, self.template_name, context)
        
        form = NewResetPasswordForm(user=context['user'], data=request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Sua senha foi alterada com sucesso!")
            return redirect(self.success_url)
        
        context['form'] = form
        return render(request, self.template_name, context)
    
    def _get_context_with_user(self, request, **kwargs):
        """
        Método auxiliar para obter o contexto com informações do usuário e validação do token
        """
        context = {
            'validlink': False,
            'user': None,
            'tenant': request.tenant,
            'tenant_name': request.tenant.schema_name,
        }
        
        try:
            # Decodificar o uidb64
            uid = force_str(urlsafe_base64_decode(kwargs.get('uidb64', '')))
            print(f"[DEBUG] UID decodificado: {uid}")
            
            try:
                # Obtenha o usuário do tenant atual usando UsuarioCliente
                user = UsuarioCliente.objects.get(pk=uid)
                print(f"[DEBUG] Usuário encontrado: {user.email}")
                
                # Verificar o token
                token = kwargs.get('token', '')
                if default_token_generator.check_token(user, token):
                    context['validlink'] = True
                    context['user'] = user
                    print(f"[DEBUG] Token válido para usuário {user.email}")
                else:
                    print(f"[DEBUG] Token inválido para usuário {user.email}")
            except UsuarioCliente.DoesNotExist:
                print(f"[DEBUG] Usuário com ID {uid} não encontrado no tenant {request.tenant.schema_name}")
        except (TypeError, ValueError, OverflowError) as e:
            print(f"[DEBUG] Erro ao decodificar UID: {str(e)}")
        
        return context


class ClienteUserMangeView(LoginRequiredMixin, View):
    template_name = 'cliente_userManage.html'
    success_url = reverse_lazy('LSCliente:CLUserList')
    
    def get_object(self):
        user_id = self.request.GET.get('id')
        if user_id:
            return get_object_or_404(UsuarioCliente, id=user_id)
        return None
    
    def get(self, request, *args, **kwargs):
        usuario = self.get_object()
        form = UsuarioClienteForm(instance=usuario) if usuario else UsuarioClienteForm()
        return render(request, self.template_name, {'form': form})
    
    def post(self, request, *args, **kwargs):
        usuario_id = request.POST.get('usuario_id')
        
        if usuario_id:
            # Modo de edição
            try:
                usuario = get_object_or_404(UsuarioCliente, id=usuario_id)
                form = UsuarioClienteForm(request.POST, request.FILES, instance=usuario)
            except Exception as e:
                messages.error(request, f'Usuário não encontrado. Erro: {str(e)}')
                return redirect(self.success_url)
        else:
            # Modo de criação
            form = UsuarioClienteForm(request.POST, request.FILES)
        
        if form.is_valid():
            try:
                form.save()
                messages.success(request, 'Usuário salvo com sucesso!')
                return redirect(self.success_url)
            except Exception as e:
                messages.error(request, f'Erro ao salvar usuário: {str(e)}')
        
        # Se o formulário tem erros ou ocorreu uma exceção
        return render(request, self.template_name, {'form': form})

class ClienteUserListView(LoginRequiredMixin, ListView):
    model = UsuarioCliente
    template_name = 'cliente_userList.html'
    context_object_name = 'usuarios'
    paginate_by = 10  # Paginação de 10 usuários por página
    
    def get_queryset(self):
        queryset = UsuarioCliente.objects.all().order_by('-date_joined')
        
        # Filtro de pesquisa
        search_query = self.request.GET.get('search')
        if search_query:
            queryset = queryset.filter(
                Q(nome__icontains=search_query) | 
                Q(email__icontains=search_query)
            )
            
        return queryset

class ClienteUserToggleStatusView(LoginRequiredMixin, View):
    def post(self, request, pk):
        try:
            usuario = get_object_or_404(UsuarioCliente, id=pk)
            usuario.is_active = not usuario.is_active
            usuario.save()
            
            if usuario.is_active:
                messages.success(request, f'Usuário {usuario.nome} ativado com sucesso!')
            else:
                messages.success(request, f'Usuário {usuario.nome} inativado com sucesso!')
                
        except Exception as e:
            messages.error(request, f'Erro ao alterar status do usuário: {str(e)}')
            
        return redirect('LSCliente:CLUserList')

@method_decorator(login_required, name='dispatch')
class ClienteProvaCreateView(LoginRequiredMixin, View):
    template_name = 'cliente_provaCreate.html' 

    def get(self, request):
        """Exibe a tela de criação de prova"""
        # Busca questões disponíveis do dashboard
        questoes_dashboard = QuestaoService.buscar_questoes_dashboard()
        
        # Busca configurações do cliente
        config_cliente = ClienteSystemSettings.obter_configuracao()
        
        context = {
            'questoes_dashboard': questoes_dashboard,
            'tipos_prova': Prova.TIPO_PROVA_CHOICES,
            'config_cliente': config_cliente
        }
        
        return render(request, self.template_name, context)
    
    def post(self, request):
        """Processa a criação da prova"""
        action = request.POST.get('action')
        
        if action == 'salvar_prova':
            return self._salvar_prova(request)
        elif action == 'exportar_pdf':
            return self._exportar_pdf(request)
        elif action == 'exportar_doc':
            return self._exportar_doc(request)
        elif action == 'buscar_questao':
            return self._buscar_questao_ajax(request)
        elif action == 'criar_questao':
            return self._criar_nova_questao(request)
        
        return JsonResponse({'error': 'Ação não reconhecida'}, status=400)
    
    def _salvar_prova(self, request):
        """Salva a prova no banco de dados"""
        try:
            titulo = request.POST.get('titulo')
            materia = request.POST.get('materia')
            tipo_prova = request.POST.get('tipo_prova')
            questoes_selecionadas = json.loads(request.POST.get('questoes_selecionadas', '[]'))
            
            # Validações básicas
            if not titulo or not materia or not tipo_prova:
                return JsonResponse({'error': 'Todos os campos são obrigatórios'}, status=400)
            
            if not questoes_selecionadas:
                return JsonResponse({'error': 'Adicione pelo menos uma questão'}, status=400)
            
            # Cria a prova
            prova = Prova.objects.create(
                titulo=titulo,
                materia=materia,
                tipo_prova=tipo_prova,
                criado_por=request.user
            )
            
            # Adiciona questões à prova
            for i, questao_data in enumerate(questoes_selecionadas, 1):
                questao_completa = QuestaoService.buscar_questao_completa(questao_data['id'])
                
                QuestaoProva.objects.create(
                    prova=prova,
                    questao_id=questao_data['id'],
                    questao_dados=questao_completa,
                    ordem=i
                )
            
            messages.success(request, 'Prova salva com sucesso!')
            return JsonResponse({'success': True, 'prova_id': prova.id})
            
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Dados das questões inválidos'}, status=400)
        except Exception as e:
            return JsonResponse({'error': f'Erro interno: {str(e)}'}, status=500)
    
    def _exportar_pdf(self, request):
        """Exporta a prova como PDF"""
        try:
            # Coleta dados do formulário
            titulo = request.POST.get('titulo', '').strip()
            materia = request.POST.get('materia', '').strip()
            tipo_prova = request.POST.get('tipo_prova', '').strip()
            questoes_data = request.POST.get('questoes_selecionadas', '[]')
            
            # Validações
            if not titulo or not materia or not tipo_prova:
                return JsonResponse({'error': 'Preencha todos os campos obrigatórios'}, status=400)
            
            try:
                questoes_selecionadas = json.loads(questoes_data)
            except json.JSONDecodeError:
                return JsonResponse({'error': 'Dados das questões inválidos'}, status=400)
            
            if not questoes_selecionadas:
                return JsonResponse({'error': 'Adicione pelo menos uma questão'}, status=400)
            
            # Cria o PDF
            buffer = BytesIO()
            doc = SimpleDocTemplate(
                buffer, 
                pagesize=A4,
                topMargin=1*inch,
                bottomMargin=1*inch,
                leftMargin=0.75*inch,
                rightMargin=0.75*inch
            )
            
            # Estilos
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            normal_style = styles['Normal']
            normal_style.fontSize = 12
            normal_style.spaceAfter = 12
            
            story = []
            
            # Busca configurações do cliente
            try:
                config_cliente = ClienteSystemSettings.obter_configuracao()
            except:
                config_cliente = None
            
            # Logo do cliente (se existir)
            if config_cliente and hasattr(config_cliente, 'imagem_home_1') and config_cliente.imagem_home_1:
                try:
                    if hasattr(config_cliente.imagem_home_1, 'path'):
                        logo_path = config_cliente.imagem_home_1.path
                        if os.path.exists(logo_path):
                            logo = Image(logo_path, width=2*inch, height=1*inch)
                            story.append(logo)
                            story.append(Spacer(1, 12))
                except Exception as e:
                    print(f"Erro ao carregar logo: {e}")
            
            # Cabeçalho da prova
            story.append(Paragraph(titulo, title_style))
            
            # Informações da prova
            info_prova = f"""
            <para align="center"><b>Matéria:</b> {materia}</para>
            <para align="center"><b>Tipo:</b> {dict(Prova.TIPO_PROVA_CHOICES).get(tipo_prova, tipo_prova)}</para>
            """
            story.append(Paragraph(info_prova, normal_style))
            story.append(Spacer(1, 20))
            
            # Campo para nome do aluno
            nome_campo = """
            <para><b>Nome:</b> _______________________________________________</para>
            <para><b>Data:</b> ___/___/______</para>
            """
            story.append(Paragraph(nome_campo, normal_style))
            story.append(Spacer(1, 30))
            
            # Linha separadora
            story.append(Paragraph("_" * 80, normal_style))
            story.append(Spacer(1, 20))
            
            # Questões
            for i, questao_data in enumerate(questoes_selecionadas, 1):
                try:
                    # Se questao_data é um dicionário completo, usa diretamente
                    if isinstance(questao_data, dict) and 'titulo' in questao_data:
                        questao = questao_data
                    else:
                        # Caso contrário, busca os dados completos
                        questao_id = questao_data.get('id') if isinstance(questao_data, dict) else questao_data
                        questao = QuestaoService.buscar_questao_completa(questao_id)
                    
                    if not questao:
                        continue
                    
                    # Número e enunciado da questão
                    titulo_questao = questao.get('titulo', f'Questão {i}')
                    story.append(Paragraph(f"<b>{i}. {titulo_questao}</b>", normal_style))
                    story.append(Spacer(1, 10))
                    
                    # Tratamento por tipo de questão
                    tipo_questao = questao.get('tipo', 'aberta')
                    
                    if tipo_questao == 'multipla':
                        alternativas = questao.get('alternativas', [])
                        for j, alt in enumerate(alternativas):
                            letra = chr(ord('a') + j)
                            texto_alt = alt.get('texto', '') if isinstance(alt, dict) else str(alt)
                            story.append(Paragraph(f"<b>{letra})</b> {texto_alt}", normal_style))
                    
                    elif tipo_questao == 'vf':
                        frases = questao.get('frases_vf', [])
                        for frase in frases:
                            texto_frase = frase.get('texto', '') if isinstance(frase, dict) else str(frase)
                            story.append(Paragraph(f"( ) {texto_frase}", normal_style))
                    
                    elif tipo_questao == 'aberta':
                        # Adiciona linhas para resposta
                        for _ in range(4):
                            story.append(Spacer(1, 20))
                            story.append(Paragraph("_" * 80, normal_style))
                    
                    story.append(Spacer(1, 25))
                    
                except Exception as e:
                    print(f"Erro ao processar questão {i}: {e}")
                    continue
            
            # Gera o PDF
            doc.build(story)
            pdf_data = buffer.getvalue()
            buffer.close()
            
            # Sanitiza o nome do arquivo
            nome_arquivo = "".join(c for c in titulo if c.isalnum() or c in (' ', '-', '_')).rstrip()
            if not nome_arquivo:
                nome_arquivo = "prova"
            
            # Retorna o PDF
            response = HttpResponse(pdf_data, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{nome_arquivo}.pdf"'
            return response
            
        except Exception as e:
            print(f"Erro ao exportar PDF: {e}")
            return JsonResponse({'error': f'Erro ao gerar PDF: {str(e)}'}, status=500)
    
    def _exportar_doc(self, request):
        """Exporta a prova como DOC"""
        if not HAS_DOCX:
            return JsonResponse({'error': 'Biblioteca python-docx não instalada'}, status=500)
        
        try:
            # Coleta dados
            titulo = request.POST.get('titulo', '').strip()
            materia = request.POST.get('materia', '').strip()
            tipo_prova = request.POST.get('tipo_prova', '').strip()
            
            # Validações
            if not titulo or not materia or not tipo_prova:
                return JsonResponse({'error': 'Preencha todos os campos obrigatórios'}, status=400)
            
            try:
                questoes_selecionadas = json.loads(request.POST.get('questoes_selecionadas', '[]'))
            except json.JSONDecodeError:
                return JsonResponse({'error': 'Dados das questões inválidos'}, status=400)
            
            if not questoes_selecionadas:
                return JsonResponse({'error': 'Adicione pelo menos uma questão'}, status=400)
            
            # Cria documento Word
            doc = Document()
            
            # Cabeçalho
            doc.add_heading(titulo, 0)
            doc.add_paragraph(f'Matéria: {materia}')
            doc.add_paragraph(f'Tipo: {dict(Prova.TIPO_PROVA_CHOICES).get(tipo_prova, tipo_prova)}')
            doc.add_paragraph()
            
            # Campo para nome
            doc.add_paragraph('Nome: _' + '_' * 50)
            doc.add_paragraph('Data: ___/___/______')
            doc.add_paragraph()
            doc.add_paragraph('_' * 80)
            doc.add_paragraph()
            
            # Questões
            for i, questao_data in enumerate(questoes_selecionadas, 1):
                try:
                    # Trata os dados da questão
                    if isinstance(questao_data, dict) and 'titulo' in questao_data:
                        questao = questao_data
                    else:
                        questao_id = questao_data.get('id') if isinstance(questao_data, dict) else questao_data
                        questao = QuestaoService.buscar_questao_completa(questao_id)
                    
                    if not questao:
                        continue
                    
                    # Enunciado
                    p = doc.add_paragraph()
                    p.add_run(f'{i}. ').bold = True
                    p.add_run(questao.get('titulo', f'Questão {i}'))
                    
                    # Conteúdo por tipo
                    tipo_questao = questao.get('tipo', 'aberta')
                    
                    if tipo_questao == 'multipla':
                        alternativas = questao.get('alternativas', [])
                        for j, alt in enumerate(alternativas):
                            letra = chr(ord('a') + j)
                            texto_alt = alt.get('texto', '') if isinstance(alt, dict) else str(alt)
                            doc.add_paragraph(f'{letra}) {texto_alt}')
                    
                    elif tipo_questao == 'vf':
                        frases = questao.get('frases_vf', [])
                        for frase in frases:
                            texto_frase = frase.get('texto', '') if isinstance(frase, dict) else str(frase)
                            doc.add_paragraph(f'( ) {texto_frase}')
                    
                    elif tipo_questao == 'aberta':
                        for _ in range(4):
                            doc.add_paragraph('_' * 80)
                    
                    doc.add_paragraph()
                    
                except Exception as e:
                    print(f"Erro ao processar questão {i}: {e}")
                    continue
            
            # Salva em buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Sanitiza nome do arquivo
            nome_arquivo = "".join(c for c in titulo if c.isalnum() or c in (' ', '-', '_')).rstrip()
            if not nome_arquivo:
                nome_arquivo = "prova"
            
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{nome_arquivo}.docx"'
            return response
            
        except Exception as e:
            print(f"Erro ao exportar DOC: {e}")
            return JsonResponse({'error': f'Erro ao gerar DOC: {str(e)}'}, status=500)
    
    def _buscar_questao_ajax(self, request):
        """Busca dados completos de uma questão via AJAX"""
        try:
            questao_id = request.POST.get('questao_id')
            if not questao_id:
                return JsonResponse({'error': 'ID da questão não fornecido'}, status=400)
            
            questao = QuestaoService.buscar_questao_completa(questao_id)
            
            if questao:
                return JsonResponse({'success': True, 'questao': questao})
            else:
                return JsonResponse({'error': 'Questão não encontrada'}, status=404)
        
        except Exception as e:
            return JsonResponse({'error': f'Erro ao buscar questão: {str(e)}'}, status=500)
    
    def _criar_nova_questao(self, request):
        """Cria uma nova questão no cliente"""
        try:
            # Implementar lógica para criar questões no tenant atual
            return JsonResponse({'success': True, 'message': 'Funcionalidade em desenvolvimento'})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)


class ClienteParametroView(LoginRequiredMixin, TemplateView):
    template_name = 'cliente_parametros.html'
    
    def get(self, request):
        # Obter ou criar configuração do cliente
        configuracao = ClienteSystemSettings.obter_configuracao()
        
        # Obter sessões ativas dos usuários - CORRIGIDO
        sessoes_ativas = []
        try:
            # Calcular tempo limite para sessões ativas
            tempo_limite = timezone.now() - timezone.timedelta(minutes=configuracao.tempo_maximo_inatividade)
            
            # Buscar sessões ativas usando o modelo correto
            from LSCliente.models import SessaoUsuarioCliente
            sessoes_query = SessaoUsuarioCliente.objects.select_related('usuario').filter(
                ultima_atividade__gte=tempo_limite
            ).order_by('-ultima_atividade')
            
            # Debug: Verificar se existem sessões
            print(f"Total de sessões encontradas: {sessoes_query.count()}")
            
            for sessao in sessoes_query:
                print(f"Sessão: {sessao.usuario.nome if hasattr(sessao.usuario, 'nome') else sessao.usuario.email}")
            
            sessoes_ativas = list(sessoes_query)
            
        except Exception as e:
            print(f"Erro ao buscar sessões ativas: {e}")
            sessoes_ativas = []
        
        # Total de usuários no sistema
        total_usuarios = UsuarioCliente.objects.filter(is_active=True).count()
        
        context = {
            'configuracao': configuracao,
            'sessoes_ativas': sessoes_ativas,
            'total_usuarios': total_usuarios,
            'title': 'Configurações do Sistema',
            # Debug info
            'debug_info': {
                'total_sessoes': len(sessoes_ativas),
                'tempo_limite': configuracao.tempo_maximo_inatividade,
            }
        }
        
        return render(request, self.template_name, context)
    
    def post(self, request):
        action = request.POST.get('action')
        configuracao = ClienteSystemSettings.obter_configuracao()
        
        try:
            if action == 'logout_user':
                user_id = request.POST.get('user_id')
                try:
                    user_id = int(user_id)
                    
                    # Buscar a sessão do usuário usando o modelo correto
                    from LSCliente.models import SessaoUsuarioCliente
                    sessao = SessaoUsuarioCliente.objects.filter(usuario_id=user_id).first()
                    
                    if sessao:
                        # Nome do usuário para a mensagem
                        usuario_nome = getattr(sessao.usuario, 'nome', None) or getattr(sessao.usuario, 'email', f'ID {user_id}')
                        
                        # Encerrar a sessão Django se existir
                        if hasattr(sessao, 'chave_sessao') and sessao.chave_sessao:
                            try:
                                from django.contrib.sessions.models import Session
                                Session.objects.filter(session_key=sessao.chave_sessao).delete()
                            except Exception as e:
                                print(f"Erro ao deletar sessão Django: {e}")
                        
                        # Remover do nosso rastreamento
                        sessao.delete()
                        
                        return JsonResponse({
                            'success': True, 
                            'message': f'Usuário {usuario_nome} deslogado com sucesso!'
                        })
                    else:
                        return JsonResponse({
                            'success': False, 
                            'message': 'Usuário não encontrado ou não está logado'
                        })
                        
                except (ValueError, TypeError):
                    return JsonResponse({'success': False, 'message': 'ID de usuário inválido'})
            
            elif action == 'update_logo':
                # Debug: Verificar o que está chegando
                print(f"DEBUG - Action: {action}")
                print(f"DEBUG - request.FILES: {dict(request.FILES)}")
                print(f"DEBUG - request.POST: {dict(request.POST)}")
                print(f"DEBUG - FILES keys: {list(request.FILES.keys())}")
                
                # Processar upload da logo - verificar diferentes nomes de campo
                logo_file = None
                possible_names = ['logo', 'imagem_home_1', 'imagem', 'image', 'file', 'upload']
                
                for name in possible_names:
                    if name in request.FILES:
                        logo_file = request.FILES[name]
                        print(f"DEBUG - Arquivo encontrado com nome: {name}")
                        break
                
                if logo_file:
                    try:
                        # Verificar se é uma imagem válida
                        if not logo_file.content_type.startswith('image/'):
                            return JsonResponse({
                                'success': False, 
                                'message': f'Arquivo deve ser uma imagem. Tipo recebido: {logo_file.content_type}'
                            })
                        
                        # Salvar no campo correto (imagem_home_1 baseado no template)
                        configuracao.imagem_home_1 = logo_file
                        configuracao.save()
                        return JsonResponse({
                            'success': True, 
                            'message': 'Logo atualizada com sucesso!',
                            'logo_url': configuracao.imagem_home_1.url if configuracao.imagem_home_1 else None
                        })
                    except Exception as e:
                        print(f"Erro ao salvar logo: {e}")
                        return JsonResponse({
                            'success': False, 
                            'message': f'Erro ao salvar imagem: {str(e)}'
                        })
                else:
                    return JsonResponse({
                        'success': False, 
                        'message': f'Nenhuma imagem foi enviada. Campos disponíveis: {list(request.FILES.keys())}'
                    })
            
            elif action == 'update_colors':
                # Atualizar cores do sistema
                try:
                    primary_color = request.POST.get('primary_color', '').strip()
                    second_color = request.POST.get('second_color', '').strip()
                    
                    if primary_color:
                        configuracao.system_primary_color = primary_color
                    if second_color:
                        configuracao.system_second_color = second_color
                    
                    configuracao.save()
                    
                    return JsonResponse({
                        'success': True, 
                        'message': 'Cores atualizadas com sucesso!'
                    })
                    
                except Exception as e:
                    return JsonResponse({
                        'success': False, 
                        'message': f'Erro ao atualizar cores: {str(e)}'
                    })
            
            elif action == 'update_timeout':
                # Atualizar tempo de inatividade
                try:
                    timeout_minutes = int(request.POST.get('timeout_minutes', 0))
                    if timeout_minutes > 0:
                        configuracao.tempo_maximo_inatividade = timeout_minutes
                        configuracao.save()
                        
                        return JsonResponse({
                            'success': True, 
                            'message': f'Tempo de inatividade atualizado para {timeout_minutes} minutos!'
                        })
                    else:
                        return JsonResponse({
                            'success': False, 
                            'message': 'Tempo deve ser maior que 0'
                        })
                        
                except (ValueError, TypeError):
                    return JsonResponse({
                        'success': False, 
                        'message': 'Tempo inválido fornecido'
                    })
            
            else:
                # Ação não reconhecida
                return JsonResponse({
                    'success': False, 
                    'message': 'Ação não reconhecida'
                })
            
        except Exception as e:
            print(f"Erro no processamento: {e}")
            return JsonResponse({
                'success': False, 
                'message': f'Erro interno do servidor: {str(e)}'
            })




def custom_logout_view(request):
    logout(request)
    return redirect('clientehome')
